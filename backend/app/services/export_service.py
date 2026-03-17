"""
Export Service
สร้างไฟล์ output รูปแบบ DOCX (และอื่นๆ ที่อาจเพิ่มในอนาคต)
"""
from typing import Dict, Any

from app.config import settings


class ExportService:
    """Service for exporting translated documents to various formats"""
    
    def export_to_docx(self, doc_result: Dict[str, Any], output_path: str) -> str:
        """
        Export to Word document
        แต่ละ page = 1 section, blocks = paragraphs
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        for page_no in range(1, doc_result["num_pages"] + 1):
            if page_no not in doc_result["pages"]:
                if str(page_no) in doc_result["pages"]:
                     page_data = doc_result["pages"][str(page_no)]
                else:
                    continue
            else:
                page_data = doc_result["pages"][page_no]
            
            # เพิ่ม page header
            if page_no > 1:
                doc.add_page_break()
            
            # เพิ่ม blocks เป็น paragraphs
            for block in page_data["blocks"]:
                text = block.get("text", "")
                label = block.get("label", "text")
                
                if not text.strip():
                    continue
                
                para = doc.add_paragraph()
                run = para.add_run(text)
                
                # Style based on label
                if "header" in label.lower() or "title" in label.lower():
                    run.bold = True
                    run.font.size = Pt(14)
                elif "list" in label.lower():
                    para.style = "List Bullet"
                else:
                    run.font.size = Pt(11)
        
        doc.save(output_path)
        return output_path


# Singleton instance
export_service = ExportService()
